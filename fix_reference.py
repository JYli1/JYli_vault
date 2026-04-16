from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from lxml import etree

BLACK = RGBColor(0, 0, 0)

# ---- Step 1: Update reference template ----
doc = Document("custom-reference.docx")

# Force all styles to SimSun + black
for style in doc.styles:
    if hasattr(style, 'font') and style.font is not None:
        style.font.name = "宋体"
        style.font.color.rgb = BLACK
        # Set eastAsia font
        rpr = style.element.find(qn("w:rPr"))
        if rpr is None:
            rpr = etree.SubElement(style.element, qn("w:rPr"))
        rfonts = rpr.find(qn("w:rFonts"))
        if rfonts is None:
            rfonts = etree.SubElement(rpr, qn("w:rFonts"))
        rfonts.set(qn("w:ascii"), "宋体")
        rfonts.set(qn("w:hAnsi"), "宋体")
        rfonts.set(qn("w:eastAsia"), "宋体")

style_config = {
    "Title": {"font_size": 22, "bold": True, "alignment": WD_ALIGN_PARAGRAPH.CENTER},
    "Heading 1": {"font_size": 16, "bold": True, "space_before": 18, "space_after": 10},
    "Heading 2": {"font_size": 14, "bold": True, "space_before": 14, "space_after": 8},
    "Heading 3": {"font_size": 12, "bold": True, "space_before": 10, "space_after": 6},
    "Normal": {"font_size": 11, "space_after": 6, "line_spacing": 1.5},
    "First Paragraph": {"font_size": 11, "space_after": 6, "line_spacing": 1.5},
}

for name, cfg in style_config.items():
    try:
        style = doc.styles[name]
    except KeyError:
        continue
    pf = style.paragraph_format
    style.font.size = Pt(cfg["font_size"])
    if cfg.get("bold"):
        style.font.bold = True
    if cfg.get("alignment"):
        pf.alignment = cfg["alignment"]
    if cfg.get("space_before"):
        pf.space_before = Pt(cfg["space_before"])
    if cfg.get("space_after"):
        pf.space_after = Pt(cfg["space_after"])
    if cfg.get("line_spacing"):
        pf.line_spacing = cfg["line_spacing"]

for section in doc.sections:
    section.top_margin = Cm(2.54)
    section.bottom_margin = Cm(2.54)
    section.left_margin = Cm(3.17)
    section.right_margin = Cm(3.17)

doc.save("custom-reference.docx")
print("Reference template updated.")

# ---- Step 2: Generate docx via pandoc ----
import subprocess
subprocess.run([
    "pandoc", "123.md", "-o", "123_final.docx",
    "--reference-doc=custom-reference.docx",
    "--syntax-highlighting=tango",
    "--toc", "--number-sections", "--wrap=none",
    "-f", "markdown+tex_math_dollars+pipe_tables+raw_html"
], check=True)
print("Pandoc conversion done.")

# ---- Step 3: Post-process - force ALL runs to SimSun + black ----
doc2 = Document("123_final.docx")

def set_run_font(run):
    run.font.name = "宋体"
    run.font.color.rgb = BLACK
    rpr = run._element.find(qn("w:rPr"))
    if rpr is None:
        rpr = etree.SubElement(run._element, qn("w:rPr"))
    rfonts = rpr.find(qn("w:rFonts"))
    if rfonts is None:
        rfonts = etree.SubElement(rpr, qn("w:rFonts"))
    rfonts.set(qn("w:ascii"), "宋体")
    rfonts.set(qn("w:hAnsi"), "宋体")
    rfonts.set(qn("w:eastAsia"), "宋体")

for para in doc2.paragraphs:
    for run in para.runs:
        set_run_font(run)

for table in doc2.tables:
    for row in table.rows:
        for cell in row.cells:
            for para in cell.paragraphs:
                for run in para.runs:
                    set_run_font(run)

doc2.save("123_final.docx")
print("Post-processing done. All fonts set to SimSun, all colors set to black.")
